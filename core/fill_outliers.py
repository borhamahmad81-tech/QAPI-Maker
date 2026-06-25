"""Outliers filler.

Excel layout: a header row containing repeated 'Name_ID' / 'ID' / value blocks.
Column map (0-indexed), derived from the sheet:

  Hb block:      Name_ID=6,  value=11  (Hb_g/dl)        -> low if <10, high if >12
  Phosph block:  Name_ID=13, value=18  (Phosph_mg/dl)   -> low if <2.5, high if >5.5
  Calcium block: Name_ID=20, value=25  (Calcium)        -> low if <8.42, high if >10
  VAC block:     Name_ID=27, value=29  (Access-can)     -> vascular access outliers
  Kt/V block:    Name_ID=32, value=34  (Kt/V)           -> Kt/V < 1.4
  Albumin block: Name_ID=36, value=38  (Albumin_g/l)    -> Alb <= 3.5
  BFR block:     Name_ID=40, value=42  (Sum of BFR)     -> BFR < 350
  MAP block:     Name_ID=45, value=47  (Average of MAP) -> MAP > 105

The Name_ID cells already read 'Surname, First (1234567890)' which is exactly the
format the filled example uses, so we use them verbatim and append ': value.'.

Doc KPI rows (Categorization text) we target:
  Vascular access ... not in Target & Not excluded.   <- VAC
  MAP> 105 mmHg                                        <- MAP
  Kt/V < 1.4                                           <- Kt/V
  Phosphorus < 2.5 / Phosphorus >5.5                  <- Phosph low/high
  Ca < 8.42 / Ca > 10                                  <- Calcium low/high
  Hb < 10 / Hb > 12                                    <- Hb low/high
  Alb <= 3.5                                           <- Albumin
  (PTH rows: no Excel source -> left blank for manual entry)

Only the "Patient /Reason" column (col index 3) is filled, as 'Name_ID: value.'.
Count column and Action plan stay untouched / manual.
"""
from docx import Document
from .helpers import read_sheet, clean_text
from .docx_tools import set_cell_text, insert_row_after


# block definition: (name_id_col, value_col, classifier)
def _num(v):
    try:
        return float(v)
    except (TypeError, ValueError):
        return None


# Each block is found by matching its VALUE-column header text. The Name_ID
# column is the nearest 'Name_ID' header to the LEFT of that value column.
# This survives column shifts as long as headers keep their names.
BLOCK_SPECS = {
    "hb_low":  (["hb_g/dl", "hb g/dl", "hb_g"],          lambda x: _num(x) is not None and _num(x) < 10),
    "hb_high": (["hb_g/dl", "hb g/dl", "hb_g"],          lambda x: _num(x) is not None and _num(x) > 12),
    "ph_low":  (["phosph_mg/dl", "phosph", "phosphor"],  lambda x: _num(x) is not None and _num(x) < 2.5),
    "ph_high": (["phosph_mg/dl", "phosph", "phosphor"],  lambda x: _num(x) is not None and _num(x) > 5.5),
    "ca_low":  (["calcium"],                              lambda x: _num(x) is not None and _num(x) < 8.42),
    "ca_high": (["calcium"],                              lambda x: _num(x) is not None and _num(x) > 10),
    "vac":     (["access-can", "access", "vac"],          lambda x: True),
    "ktv":     (["kt/v"],                                 lambda x: _num(x) is not None and _num(x) < 1.4),
    "alb":     (["albumin_g/l", "albumin"],               lambda x: _num(x) is not None and (_num(x) / 10.0 if _num(x) > 10 else _num(x)) <= 3.5),
    "bfr":     (["sum of bfr", "bfr"],                    lambda x: _num(x) is not None and _num(x) < 350),
    "map":     (["average of map", "map"],                lambda x: _num(x) is not None and _num(x) > 105),
    "pth_low":  (["pth_pg/ml", "pth"],                    lambda x: _num(x) is not None and _num(x) < 150),
    "pth_high": (["pth_pg/ml", "pth"],                    lambda x: _num(x) is not None and _num(x) > 600),
}


def _find_header_row(rows):
    for i, r in enumerate(rows):
        cells = [str(c).strip().lower() if c is not None else "" for c in r]
        if cells.count("name_id") >= 2:
            return i, cells
    raise ValueError("Could not find the Outliers header row (needs 'Name_ID' labels).")


def _locate_columns(header_cells, value_aliases):
    """Find the value column matching any alias, and the nearest Name_ID column
    to its left. Returns (name_col, value_col) or None."""
    for vc, label in enumerate(header_cells):
        if any(alias == label or alias in label for alias in value_aliases):
            # nearest Name_ID to the left
            for nc in range(vc - 1, -1, -1):
                if header_cells[nc] == "name_id":
                    return nc, vc
    return None

# Match a Doc row to a block key by keywords found in the Categorization cell.
def _row_block_key(categorization, kpi):
    c = categorization.lower()
    k = kpi.lower()
    if "not in target" in c and "exclud" in c:
        return "vac"
    if "map" in c and "105" in c:
        return "map"
    if "kt/v" in c and "1.4" in c:
        return "ktv"
    if "phosphorus" in c and "< 2.5" in c:
        return "ph_low"
    if "phosphorus" in c and ("5.5" in c) and (">" in c):
        return "ph_high"
    if ("ca <" in c) or ("ca <8.42" in c) or ("< 8.42" in c):
        return "ca_low"
    if ("ca >" in c) or ("> 10" in c):
        return "ca_high"
    if "hb < 10" in c:
        return "hb_low"
    if "hb > 12" in c:
        return "hb_high"
    if "alb" in c and "3.5" in c:
        return "alb"
    if "bfr" in c or "blood flow" in c:
        return "bfr"
    if "pth" in c and "< 150" in c:
        return "pth_low"
    if "pth" in c and ("> 600" in c or "600" in c and ">" in c):
        return "pth_high"
    if "pth" in c or "pth" in k:
        return None  # PTH row but no recognizable threshold
    return None


def _collect(rows, start_row, name_col, val_col, classifier, block_key=None):
    out = []
    for r in rows[start_row + 1:]:
        if name_col >= len(r) or val_col >= len(r):
            continue
        nid = r[name_col]
        val = r[val_col]
        if nid is None:
            continue
        s = clean_text(nid)
        if not s or s.lower() in ("name_id", "(blank)"):
            continue
        if classifier(val):
            out.append((s, val, block_key))
    return out


def _fmt_value(v, block_key=None):
    if v is None:
        return ""
    # Albumin is stored as g/L but reported as g/dl (divide by 10).
    if block_key == "alb":
        n = _num(v)
        if n is not None and n > 10:
            v = n / 10.0
    if isinstance(v, float):
        if v.is_integer():
            return str(int(v))
        return f"{round(v, 3):g}"
    n = _num(v)
    if n is not None and not float(n).is_integer():
        return f"{round(n, 3):g}"
    return str(v).strip()


def _find_table(doc):
    for t in doc.tables:
        for row in t.rows:
            joined = " ".join(c.text.lower() for c in row.cells)
            if "outliers" in joined and "patient" in joined and "action plan" in joined:
                return t
    # fallback: the big 5-col table
    for t in doc.tables:
        if len(t.columns) == 5:
            return t
    raise ValueError("Outliers table not found.")


def fill(doc_path, xlsx_path, out_path):
    rows = read_sheet(xlsx_path)
    header_row, header_cells = _find_header_row(rows)

    # Pre-collect patients for each block using header-located columns
    block_data = {}
    for key, (value_aliases, classifier) in BLOCK_SPECS.items():
        loc = _locate_columns(header_cells, value_aliases)
        if loc is None:
            block_data[key] = []
            continue
        nc, vc = loc
        block_data[key] = _collect(rows, header_row, nc, vc, classifier, key)

    doc = Document(doc_path)
    table = _find_table(doc)

    # If the template lacks a BFR row but the Excel has BFR outliers, clone the
    # Kt/V row and insert a BFR row right after it (matching the filled example).
    if block_data.get("bfr"):
        has_bfr_row = any(
            "bfr" in r.cells[1].text.lower() or "blood flow" in r.cells[0].text.lower()
            for r in table.rows if len(r.cells) >= 2
        )
        if not has_bfr_row:
            ktv_idx = None
            for ri, row in enumerate(table.rows):
                if len(row.cells) >= 2 and "kt/v" in row.cells[1].text.lower():
                    ktv_idx = ri
            if ktv_idx is not None:
                new_row = insert_row_after(table, ktv_idx)
                set_cell_text(new_row.cells[0], "Blood flow rate (BFR): >350 ml/min")
                set_cell_text(new_row.cells[1], "# Pts BFR < 350 ml/min")

    # Identify KPI/Categorization columns: col0=KPI, col1=Categorization,
    # col2=Count, col3=Patient/Reason, col4=Action plan
    # Walk template rows, group consecutive rows by their categorization key.
    # For each distinct categorization block we fill rows with that block's pts.

    # First, map each template data row -> block key
    row_keys = {}
    header_done = False
    for ri, row in enumerate(table.rows):
        cells = row.cells
        if len(cells) < 5:
            continue
        kpi = cells[0].text.strip()
        cat = cells[1].text.strip()
        joined = (kpi + " " + cat).lower()
        if "categorization" in joined or "details of patients" in joined:
            header_done = True
            continue
        if "attach a report" in kpi.lower():
            continue
        if not header_done:
            continue
        key = _row_block_key(cat, kpi)
        row_keys[ri] = key

    # Group template rows by (kpi, cat) contiguous block to know how many
    # template rows each KPI currently provides.
    # We'll insert extra rows after the last row of each group as needed.
    # Build groups: list of (key, [row_indices])
    groups = []
    cur_key = object()
    cur_rows = []
    cur_sig = None
    for ri in sorted(row_keys):
        cells = table.rows[ri].cells
        sig = (cells[0].text.strip(), cells[1].text.strip())
        if sig != cur_sig:
            if cur_rows:
                groups.append((cur_key, cur_sig, cur_rows))
            cur_key = row_keys[ri]
            cur_sig = sig
            cur_rows = [ri]
        else:
            cur_rows.append(ri)
    if cur_rows:
        groups.append((cur_key, cur_sig, cur_rows))

    # Fill from the BOTTOM up so row insertions don't shift indices of groups
    # we haven't processed yet.
    total = 0
    for key, sig, idxs in reversed(groups):
        if key is None:
            continue
        pts = block_data.get(key, [])
        if not pts:
            continue
        # ensure enough rows in this group
        need = len(pts)
        have = len(idxs)
        last_idx = idxs[-1]
        while have < need:
            insert_row_after(table, last_idx + (have - len(idxs)))
            have += 1
        # re-fetch group row indices (they are contiguous starting at idxs[0])
        start = idxs[0]
        for i, (nid, val, bkey) in enumerate(pts):
            row = table.rows[start + i]
            patient_text = f"{nid}: {_fmt_value(val, bkey)}."
            set_cell_text(row.cells[3], patient_text)
            total += 1

        # Merge the Count column (index 2) across this KPI block and write the
        # number of patients actually filled.
        try:
            first_count = table.rows[start].cells[2]
            last_count = table.rows[start + len(pts) - 1].cells[2]
            merged = first_count.merge(last_count)
            set_cell_text(merged, str(len(pts)))
        except Exception:
            # if merge isn't possible (already merged differently), just write
            set_cell_text(table.rows[start].cells[2], str(len(pts)))

    doc.save(out_path)
    return total
