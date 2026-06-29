"""Compare two filled QAPI documents (e.g. program-generated vs manually
prepared) and report DIFFERENCES ONLY across all five sections.

Match keys and tolerances (locked down with the user):
  Outliers        : (ID, KPI block)         lab value, tolerance 0.1
  Excluded        : (ID, KPI category)       presence only (text ignored)
  Hospitalization : (ID, Admission date)     discharge date compared
  Extra Sessions  : (ID)                     session count, exact
  KPI Medical     : (KPI name)                % Before/After tol 0.1, excluded exact

Reason / Action Plan / Cause / Hospital / wording are never compared -- only
the keys and the specific values listed above.
"""
import re
from docx import Document

from .fill_outliers import _find_table as _find_outliers_table
from .fill_outliers import _row_block_key, _extract_id as _ol_extract_id
from .fill_excluded import _find_table as _find_excluded_table
from .fill_excluded import _row_texts as _excluded_row_texts
from .fill_excluded import _extract_id as _ex_extract_id
from .fill_excluded import _doc_category_key
from .fill_kpi_medical import _find_medical_table, _excel_key
from .helpers import clean_id

TOL = 0.1


def _num(v):
    if v is None:
        return None
    s = str(v).strip().replace("%", "")
    try:
        return float(s)
    except ValueError:
        return None


def _close(a, b, tol=TOL):
    na, nb = _num(a), _num(b)
    if na is None or nb is None:
        return (a or "").strip() == (b or "").strip()
    return abs(na - nb) <= tol


# ---------------------------------------------------------------- Outliers --
def _extract_outliers(doc):
    """Return {(id, block_key): value_text}."""
    out = {}
    try:
        table = _find_outliers_table(doc)
    except ValueError:
        return out
    last_kpi = last_cat = ""
    for row in table.rows[2:]:
        cells = row.cells
        if len(cells) < 5:
            continue
        kpi = cells[0].text.strip() or last_kpi
        cat = cells[1].text.strip() or last_cat
        if cells[0].text.strip():
            last_kpi = cells[0].text.strip()
        if cells[1].text.strip():
            last_cat = cells[1].text.strip()
        patient_text = cells[3].text.strip()
        if not patient_text or "action plan" in cells[4].text.lower():
            continue
        pid = _ol_extract_id(patient_text)
        if not pid:
            continue
        key_block = _row_block_key(cat, kpi)
        if not key_block:
            continue
        # value = first number after the ID in the patient cell
        idpos = patient_text.find(pid)
        tail = patient_text[idpos + len(pid):] if idpos != -1 else patient_text
        m = re.search(r"\d+(?:\.\d+)?", tail)
        value = m.group(0) if m else ""
        out[(pid, key_block)] = value
    return out


# --------------------------------------------------------------- Excluded ---
def _extract_excluded(doc):
    """Return set of (id, category_key) present -- presence only."""
    out = set()
    table = _find_excluded_table(doc)
    if table is None:
        return out
    last_cat_key = None
    # mirror the leading-row lookahead used by fill_excluded
    raw = []
    for ri in range(2, len(table.rows)):
        raw.append((ri,) + _excluded_row_texts(table, ri))
    first_label = None
    for ri, idt, cat, rs, pl in raw:
        if cat is not None and cat.strip():
            first_label = cat.strip()
            break
    last_cat_label = None
    for ri, idt, cat, rs, pl in raw:
        if cat is not None and cat.strip():
            last_cat_label = cat.strip()
            last_cat_key = _doc_category_key(cat)
        elif last_cat_key is None and first_label:
            last_cat_key = _doc_category_key(first_label)
        pid = _ex_extract_id(idt)
        if pid and last_cat_key:
            out.add((pid, last_cat_key))
    return out


# ------------------------------------------------------------ Hospitalization
def _extract_hospitalization(doc):
    """Return {(id, admission_date): discharge_date_or_empty}."""
    out = {}
    target = None
    for t in doc.tables:
        for row in t.rows:
            if row.cells and "hospitalized patients" in row.cells[0].text.lower():
                target = row.cells[2].text
                break
        if target:
            break
    if not target:
        return out
    # entries are separated by blank lines, "N. Name: ID" then indented fields
    blocks = re.split(r"\n\s*\n", target)
    for b in blocks:
        idm = re.search(r"(\d{10})", b)
        if not idm:
            continue
        pid = clean_id(idm.group(1))
        adm = re.search(r"Date admitted:\s*([0-9/]+)", b)
        dis = re.search(r"Discharged on\s*([0-9/]+)", b)
        adm_date = adm.group(1) if adm else ""
        dis_date = dis.group(1) if dis else ""
        out[(pid, adm_date)] = dis_date
    return out


# --------------------------------------------------------------- Extrasessions
def _extract_extrasessions(doc):
    """Return {id: session_count}."""
    out = {}
    table = None
    for t in doc.tables:
        for row in t.rows:
            joined = " ".join(c.text.lower() for c in row.cells)
            if "patient id" in joined and "extra session" in joined:
                table = t
                break
        if table:
            break
    if not table:
        return out
    for row in table.rows[2:]:
        cells = row.cells
        if len(cells) < 2:
            continue
        c0 = cells[0].text.strip()
        idm = re.search(r"(\d{10})", c0)
        if not idm:
            continue
        pid = clean_id(idm.group(1))
        c1 = cells[1].text.strip()
        m = re.search(r"(\d+)\s*session", c1)
        out[pid] = m.group(1) if m else ""
    return out


# ------------------------------------------------------------------ KPI Med --
def _extract_kpi_medical(doc):
    """Return {kpi_key: (before, after, excluded)}."""
    out = {}
    table = _find_medical_table(doc)
    if table is None:
        return out
    hdr = [c.text.strip().lower() for c in table.rows[1].cells]
    def col(*names):
        for idx, h in enumerate(hdr):
            if any(n in h for n in names):
                return idx
        return None
    c_before = col("before exclusion")
    c_after = col("after exclusion")
    c_excl = col("excluded")
    for row in table.rows[2:]:
        name = row.cells[0].text.strip()
        key = _excel_key(name)
        if not key:
            continue
        before = row.cells[c_before].text.strip() if c_before is not None else ""
        after = row.cells[c_after].text.strip() if c_after is not None else ""
        excl = row.cells[c_excl].text.strip() if c_excl is not None else ""
        out[key] = (before, after, excl)
    return out


# ----------------------------------------------------------------- Compare --
def _diff_dict(a, b, label_fn, value_label="value"):
    """Generic differ for {key: value} dicts. Returns list of strings."""
    diffs = []
    for k in sorted(set(a) | set(b), key=lambda x: str(x)):
        if k in a and k not in b:
            diffs.append(f"Only in Doc A: {label_fn(k)} ({value_label}={a[k]!r})")
        elif k in b and k not in a:
            diffs.append(f"Only in Doc B: {label_fn(k)} ({value_label}={b[k]!r})")
        elif not _close(a[k], b[k]):
            diffs.append(f"Mismatch: {label_fn(k)} -> A={a[k]!r} vs B={b[k]!r}")
    return diffs


def _diff_set(a, b, label_fn):
    diffs = []
    for k in sorted((a | b), key=lambda x: str(x)):
        if k in a and k not in b:
            diffs.append(f"Only in Doc A: {label_fn(k)}")
        elif k in b and k not in a:
            diffs.append(f"Only in Doc B: {label_fn(k)}")
    return diffs


def compare(doc_a_path, doc_b_path):
    """Compare two filled QAPI documents. Returns {section: [diff strings]}."""
    a = Document(doc_a_path)
    b = Document(doc_b_path)
    report = {}

    oa, ob = _extract_outliers(a), _extract_outliers(b)
    report["Outliers"] = _diff_dict(
        oa, ob, lambda k: f"ID {k[0]} / {k[1]}", "lab value")

    ea, eb = _extract_excluded(a), _extract_excluded(b)
    report["Excluded Patients"] = _diff_set(
        ea, eb, lambda k: f"ID {k[0]} / {k[1]}")

    ha, hb = _extract_hospitalization(a), _extract_hospitalization(b)
    report["Hospitalization"] = _diff_dict(
        ha, hb, lambda k: f"ID {k[0]} admitted {k[1]}", "discharge date")

    xa, xb = _extract_extrasessions(a), _extract_extrasessions(b)
    report["Extra Sessions"] = _diff_dict(
        xa, xb, lambda k: f"ID {k}", "session count")

    ka, kb = _extract_kpi_medical(a), _extract_kpi_medical(b)
    kpi_diffs = []
    for key in sorted(set(ka) | set(kb)):
        if key not in ka:
            kpi_diffs.append(f"Only in Doc B: KPI {key}")
            continue
        if key not in kb:
            kpi_diffs.append(f"Only in Doc A: KPI {key}")
            continue
        ba, aa, xa_ = ka[key]
        bb, ab, xb_ = kb[key]
        if not _close(ba, bb):
            kpi_diffs.append(f"Mismatch: KPI {key} % Before -> A={ba!r} vs B={bb!r}")
        if not _close(aa, ab):
            kpi_diffs.append(f"Mismatch: KPI {key} % After -> A={aa!r} vs B={ab!r}")
        if not _close(xa_, xb_, tol=0):
            kpi_diffs.append(f"Mismatch: KPI {key} No. excluded -> A={xa_!r} vs B={xb_!r}")
    report["KPI (Medical)"] = kpi_diffs

    return report


def format_report(report):
    """Render the {section: [diffs]} report as readable text."""
    lines = []
    total = 0
    for section, diffs in report.items():
        if not diffs:
            continue
        lines.append(f"=== {section} ({len(diffs)} difference(s)) ===")
        for d in diffs:
            lines.append(f"  - {d}")
        lines.append("")
        total += len(diffs)
    if total == 0:
        return "No differences found across all five sections."
    return "\n".join(lines).strip()
