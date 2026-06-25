"""Low-level docx helpers that PRESERVE formatting.

Strategy:
- To write text into a cell without disturbing fonts, we keep the first run's
  formatting and replace only its text, deleting extra runs.
- To add table rows, we deep-copy an existing data row's XML (carrying borders,
  shading, cell widths, font of the template) and then overwrite the text.
"""
import copy
from docx.oxml.ns import qn


def _first_paragraph(cell):
    return cell.paragraphs[0]


def set_cell_text(cell, text, keep_font_from_run=True):
    """Replace a cell's entire text with `text`, preserving the font of the
    first existing run. Supports multi-line text via line breaks."""
    text = "" if text is None else str(text)
    para = _first_paragraph(cell)

    # Remove all paragraphs after the first (keep formatting anchor in p[0])
    for extra in cell.paragraphs[1:]:
        extra._element.getparent().remove(extra._element)

    runs = para.runs
    if runs:
        template_run = runs[0]
        # delete trailing runs, keep the first as our formatting template
        for r in runs[1:]:
            r._element.getparent().remove(r._element)
    else:
        template_run = para.add_run("")

    # Write text into template_run, converting \n into <w:br/>
    lines = text.split("\n")
    template_run.text = lines[0]
    for line in lines[1:]:
        br = template_run._element.makeelement(qn("w:br"), {})
        template_run._element.append(br)
        t = template_run._element.makeelement(qn("w:t"), {})
        t.set(qn("xml:space"), "preserve")
        t.text = line
        template_run._element.append(t)


def clone_row(table, src_row_idx):
    """Deep-copy the row at src_row_idx, append it to the table, return new row."""
    src_tr = table.rows[src_row_idx]._tr
    new_tr = copy.deepcopy(src_tr)
    src_tr.addnext(new_tr)
    # python-docx rebuilds .rows lazily from XML; fetch the newly added one
    return table.rows[src_row_idx + 1]


def insert_row_after(table, src_row_idx):
    """Clone formatting of src row, insert directly after it, blank the text."""
    new_row = clone_row(table, src_row_idx)
    for cell in new_row.cells:
        set_cell_text(cell, "")
    return new_row


def delete_row(table, row_idx):
    """Remove the row at row_idx from the table."""
    tr = table.rows[row_idx]._tr
    tr.getparent().remove(tr)


def _get_tcPr(tc):
    tcPr = tc.find(qn("w:tcPr"))
    if tcPr is None:
        tcPr = tc.makeelement(qn("w:tcPr"), {})
        tc.insert(0, tcPr)
    return tcPr


def _set_vmerge(tc, val):
    """Set w:vMerge on a cell. val='restart' or 'continue'.

    tcPr ordering matters in OOXML; vMerge should sit early in tcPr. We remove
    any existing vMerge and insert the new one as the first child of tcPr.
    """
    tcPr = _get_tcPr(tc)
    for vm in tcPr.findall(qn("w:vMerge")):
        tcPr.remove(vm)
    vm = tcPr.makeelement(qn("w:vMerge"), {})
    vm.set(qn("w:val"), "restart" if val == "restart" else "continue")
    tcPr.insert(0, vm)


def _row_tcs(table, row_idx):
    """Return the raw <w:tc> elements of a row by grid position, WITHOUT going
    through python-docx's .cells (which traverses vMerge chains and can crash
    on partially-merged tables)."""
    tr = table.rows[row_idx]._tr
    return tr.findall(qn("w:tc"))


def _tc_set_text(tc, text):
    """Set text on a raw <w:tc>, preserving the first run's formatting."""
    from docx.table import _Cell
    cell = _Cell(tc, None)
    set_cell_text(cell, text)


def vmerge_column(table, col_idx, start_row, end_row, text=None):
    """Vertically merge a column from start_row..end_row (inclusive) using raw
    <w:tc> access so it is safe on tables that already contain other merges.

    Top cell = 'restart' carrying the text; lower cells = 'continue', each kept
    with a valid empty paragraph.
    """
    if end_row < start_row:
        return
    top_tcs = _row_tcs(table, start_row)
    if col_idx >= len(top_tcs):
        return
    if text is not None:
        _tc_set_text(top_tcs[col_idx], text)
    if end_row == start_row:
        return
    _set_vmerge(top_tcs[col_idx], "restart")
    for r in range(start_row + 1, end_row + 1):
        tcs = _row_tcs(table, r)
        if col_idx < len(tcs):
            _tc_set_text(tcs[col_idx], "")
            _set_vmerge(tcs[col_idx], "continue")
