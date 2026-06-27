"""Outliers filler.

Excel layout: a header row containing repeated 'Name_ID' / 'ID' / value blocks.
Column map (0-indexed), derived from the sheet:

  Hb block:      Name_ID=6,  value=11  (Hb_g/dl)        -> low if <10, high if >12
  Phosph block:  Name_ID=13, value=18  (Phosph_mg/dl)   -> low if <2.5, high if >5.5
  Calcium block: Name_ID=20, value=25  (Calcium)        -> low if <8.42, high if >10
  VAC block:     Name_ID=27, value=29  (Access-can)     -> vascular access outliers
  Kt/V block:    Name_ID=32, value=34  (Kt/V)           -> Kt/V < 1.4
  Albumin block: Name_ID=36, value=38  (Albumin_g/l)    -> Alb <= 3.5
  BFR block:     Name_ID=40, value=42  (Sum of BFR)     -> BFR < 350
  MAP block:     Name_ID=45, value=47  (Average of MAP) -> MAP > 105

The Name_ID cells already read 'Surname, First (1234567890)' which is exactly the
format the filled example uses, so we use them verbatim and append ': value.'.

Doc KPI rows (Categorization text) we target:
  Vascular access ... not in Target & Not excluded.   <- VAC
  MAP> 105 mmHg                                        <- MAP
  Kt/V < 1.4                                           <- Kt/V
  Phosphorus < 2.5 / Phosphorus >5.5                  <- Phosph low/high
  Ca < 8.42 / Ca > 10                                  <- Calcium low/high
  Hb < 10 / Hb > 12                                    <- Hb low/high
  Alb <= 3.5                                           <- Albumin
  (PTH rows: no Excel source -> left blank for manual entry)

Only the "Patient /Reason" column (col index 3) is filled, as 'Name_ID: value.'.
Count column and Action plan stay untouched / manual.
"""
from docx import Document
from .helpers import read_sheet, clean_text, clean_id
from .docx_tools import set_cell_text, insert_row_after, delete_row, vmerge_column

import re

_ID_RE = re.compile(r"\b(\d{10})\b")
# A numeric value token: integer or decimal, e.g. 9.6, 113, 2500, 10.74
_VALUE_RE = re.compile(r"\d+(?:\.\d+)?")


def _extract_id(text):
    """Return the first 10-digit ID found in a cell, MOH-stripped, or ''."""
    if not text:
        return ""
    # Prefer an ID inside parentheses if present.
    m = re.search(r"\((\D*?)(\d{10})\D*?\)", text)
    if m:
        return clean_id(m.group(2))
    m = _ID_RE.search(text)
    return clean_id(m.group(1)) if m else ""


def _replace_value_after_id(cell_text, new_value, id_str):
    """Replace the first numeric token that appears AFTER the patient's ID in
    the cell with new_value, preserving name, ID and any trailing reason text.

    If no numeric token is found after the ID, the text is returned unchanged
    (used for text-only KPIs / cells the user keeps as-is).
    """
    if not cell_text:
        return cell_text
    # Find where the ID ends so we only touch numbers after it.
    id_pos = cell_text.find(id_str)
    if id_pos == -1:
        # ID may be formatted differently; fall back to first 10-digit run
        m = _ID_RE.search(cell_text)
        id_pos = m.start() if m else 0
        id_len = 10 if m else 0
    else:
        id_len = len(id_str)
    after_start = id_pos + id_len
    head = cell_text[:after_start]
    tail = cell_text[after_start:]
    # Replace the FIRST numeric token in the tail (the lab value).
    m = _VALUE_RE.search(tail)
    if not m:
        return cell_text  # nothing numeric to update
    new_tail = tail[:m.start()] + _fmt_value_plain(new_value) + tail[m.end():]
    return head + new_tail


def _fmt_value_plain(v):
    """Format a value as a plain string (no trailing period)."""
    if v is None:
        return ""
    if isinstance(v, float):
        if v.is_integer():
            return str(int(v))
        return f"{round(v, 3):g}"
    return str(v).strip()


# block definition: (name_id_col, value_col, classifier)
def _num(v):
    try:
        return float(v)
    except (TypeError, ValueError):
        return None


# Each block is found by matching its VALUE-column header text. The Name_ID
# column is the nearest 'Name_ID' header to the LEFT of that value column.
# This survives column shifts as long as headers keep their names.
BLOCK_SPECS = {
    "hb_low":  (["hb_g/dl", "hb g/dl", "hb_g"],          lambda x: _num(x) is not None and _num(x) < 10),
    "hb_high": (["hb_g/dl", "hb g/dl", "hb_g"],          lambda x: _num(x) is not None and _num(x) > 12),
    "ph_low":  (["phosph_mg/dl", "phosph", "phosphor"],  lambda x: _num(x) is not None and _num(x) < 2.5),
    "ph_high": (["phosph_mg/dl", "phosph", "phosphor"],  lambda x: _num(x) is not None and _num(x) > 5.5),
    "ca_low":  (["calcium"],                              lambda x: _num(x) is not None and _num(x) < 8.42),
    "ca_high": (["calcium"],                              lambda x: _num(x) is not None and _num(x) > 10),
    "vac":     (["access-can", "access", "vac"],          lambda x: True),
    "ktv":     (["kt/v"],                                 lambda x: _num(x) is not None and _num(x) < 1.4),
    "alb":     (["albumin_g/l", "albumin"],               lambda x: _num(x) is not None and (_num(x) / 10.0 if _num(x) > 10 else _num(x)) <= 3.5),
    "bfr":     (["sum of bfr", "bfr"],                    lambda x: _num(x) is not None and _num(x) < 350),
    "map":     (["average of map", "map"],                lambda x: _num(x) is not None and _num(x) > 105),
    "pth_low":  (["pth_pg/ml", "pth"],                    lambda x: _num(x) is not None and _num(x) < 150),
    "pth_high": (["pth_pg/ml", "pth"],                    lambda x: _num(x) is not None and _num(x) > 600),
}


def _find_header_row(rows):
    for i, r in enumerate(rows):
        cells = [str(c).strip().lower() if c is not None else "" for c in r]
        if cells.count("name_id") >= 2:
            return i, cells
    raise ValueError("Could not find the Outliers header row (needs 'Name_ID' labels).")


def _locate_columns(header_cells, value_aliases):
    """Find the value column matching any alias, and the nearest Name_ID column
    to its left. Returns (name_col, value_col) or None."""
    for vc, label in enumerate(header_cells):
        if any(alias == label or alias in label for alias in value_aliases):
            # nearest Name_ID to the left
            for nc in range(vc - 1, -1, -1):
                if header_cells[nc] == "name_id":
                    return nc, vc
    return None


def _locate_exclusion_col(header_cells, name_col, value_col):
    """Find this block's exclusion column: a header containing 'ex_code' or
    'exclusion' sitting BETWEEN the Name_ID column and the value column (that's
    where DaVita places HB_ex_code, BMD_ex_code, VAC_exclusion). Returns the
    column index or None."""
    lo = min(name_col, value_col)
    hi = max(name_col, value_col)
    for j in range(lo, hi + 1):
        h = header_cells[j]
        if "ex_code" in h or "exclusion" in h:
            return j
    return None


def _is_excluded(cell):
    """True if an exclusion-column cell marks the patient as excluded."""
    if cell is None:
        return False
    s = str(cell).strip()
    if not s or s.lower() in ("(blank)", "none"):
        return False
    return True

# Match a Doc row to a block key by keywords found in the Categorization cell.
def _row_block_key(categorization, kpi):
    c = categorization.lower()
    k = kpi.lower()
    if "not in target" in c and "exclud" in c:
        return "vac"
    if "map" in c and "105" in c:
        return "map"
    if "kt/v" in c and "1.4" in c:
        return "ktv"
    if "phosphorus" in c and "< 2.5" in c:
        return "ph_low"
    if "phosphorus" in c and ("5.5" in c) and (">" in c):
        return "ph_high"
    if ("ca <" in c) or ("ca <8.42" in c) or ("< 8.42" in c):
        return "ca_low"
    if ("ca >" in c) or ("> 10" in c):
        return "ca_high"
    if "hb < 10" in c:
        return "hb_low"
    if "hb > 12" in c:
        return "hb_high"
    if "alb" in c and "3.5" in c:
        return "alb"
    if "bfr" in c or "blood flow" in c:
        return "bfr"
    if "pth" in c and "< 150" in c:
        return "pth_low"
    if "pth" in c and ("> 600" in c or "600" in c and ">" in c):
        return "pth_high"
    if "pth" in c or "pth" in k:
        return None  # PTH row but no recognizable threshold
    return None


def _collect(rows, start_row, name_col, val_col, classifier, block_key=None,
             excl_col=None):
    out = []
    for r in rows[start_row + 1:]:
        if name_col >= len(r) or val_col >= len(r):
            continue
        nid = r[name_col]
        val = r[val_col]
        if nid is None:
            continue
        s = clean_text(nid)
        if not s or s.lower() in ("name_id", "(blank)"):
            continue
        # Skip patients who are EXCLUDED for this KPI — they belong in the
        # "Details of Excluded Patients" section, not the outliers list.
        if excl_col is not None and excl_col < len(r) and _is_excluded(r[excl_col]):
            continue
        if classifier(val):
            out.append((s, val, block_key))
    return out


def _fmt_value(v, block_key=None):
    if v is None:
        return ""
    # Albumin is stored as g/L but reported as g/dl (divide by 10).
    if block_key == "alb":
        n = _num(v)
        if n is not None and n > 10:
            v = n / 10.0
    if isinstance(v, float):
        if v.is_integer():
            return str(int(v))
        return f"{round(v, 3):g}"
    n = _num(v)
    if n is not None and not float(n).is_integer():
        return f"{round(n, 3):g}"
    return str(v).strip()


def _find_table(doc):
    for t in doc.tables:
        whole = " ".join(c.text.lower() for row in t.rows for c in row.cells)
        # The Outliers table is identified by its title and its column headers.
        if ("not meeting target" in whole or "outliers" in whole) \
                and "patient /reason" in whole.replace("  ", " "):
            return t
    # Secondary: title match alone.
    for t in doc.tables:
        if t.rows and "not meeting target" in t.rows[0].cells[0].text.lower():
            return t
    raise ValueError("Outliers table not found.")


def fill_doc(doc, xlsx_path):
    """Reconcile the Outliers table inside an already-open document."""
    rows = read_sheet(xlsx_path)
    header_row, header_cells = _find_header_row(rows)

    # Pre-collect patients for each block using header-located columns.
    # For each block, also locate its exclusion column so excluded patients are
    # filtered out of the outliers (they go in the Excluded section instead).
    block_data = {}
    for key, (value_aliases, classifier) in BLOCK_SPECS.items():
        loc = _locate_columns(header_cells, value_aliases)
        if loc is None:
            block_data[key] = []
            continue
        nc, vc = loc
        excl_col = _locate_exclusion_col(header_cells, nc, vc)
        # The vascular block's "value" header is 'Access-can'; its exclusion
        # marker is the adjacent 'VAC_exclusion' column. _locate_exclusion_col
        # finds it between Name_ID and the value column. If the exclusion column
        # happens to sit just AFTER the value column (vac layout), search right.
        if excl_col is None:
            for j in range(vc, min(vc + 3, len(header_cells))):
                if "exclusion" in header_cells[j] or "ex_code" in header_cells[j]:
                    excl_col = j
                    break
        block_data[key] = _collect(rows, header_row, nc, vc, classifier, key,
                                   excl_col=excl_col)

    table = _find_table(doc)

    # If the template lacks a BFR row but the Excel has BFR outliers, clone the
    # Kt/V row and insert a BFR row right after it (matching the filled example).
    if block_data.get("bfr"):
        has_bfr_row = any(
            "bfr" in r.cells[1].text.lower() or "blood flow" in r.cells[0].text.lower()
            for r in table.rows if len(r.cells) >= 2
        )
        if not has_bfr_row:
            ktv_idx = None
            for ri, row in enumerate(table.rows):
                if len(row.cells) >= 2 and "kt/v" in row.cells[1].text.lower():
                    ktv_idx = ri
            if ktv_idx is not None:
                new_row = insert_row_after(table, ktv_idx)
                set_cell_text(new_row.cells[0], "Blood flow rate (BFR): >350 ml/min")
                set_cell_text(new_row.cells[1], "# Pts BFR < 350 ml/min")

    # Identify columns: col0=KPI, col1=Categorization, col2=Count,
    # col3=Patient/Reason, col4=Action plan.
    #
    # Strategy:
    #  1. Find the data rows (between the header and the "Attach a report" row).
    #  2. For each data row capture its ORIGINAL KPI label and Categorization
    #     label, plus the block key. Cloned rows must inherit these so labels are
    #     never lost when a threshold expands to several patient rows.
    #  3. Fill patients per Categorization group, cloning rows as needed and
    #     re-stamping the KPI + Categorization text on every row.
    #  4. Merge the Count column per Categorization group (one number).
    #  5. Merge the Categorization column per group, and the KPI column per
    #     KPI-item (a KPI item can span several Categorization groups, e.g.
    #     Phosphorus low + high).

    header_done = False
    data = []  # list of dicts per ORIGINAL data row
    for ri, row in enumerate(table.rows):
        cells = row.cells
        if len(cells) < 5:
            continue
        kpi = cells[0].text.strip()
        cat = cells[1].text.strip()
        joined = (kpi + " " + cat).lower()
        if "categorization" in joined or "details of patients" in joined:
            header_done = True
            continue
        if "attach a report" in kpi.lower():
            continue
        if not header_done:
            continue
        # Capture any prefilled patient text + action plan on this row.
        patient_text = cells[3].text.strip()
        action_plan = cells[4].text.strip()
        pid = _extract_id(patient_text)
        data.append({"idx": ri, "kpi": kpi, "cat": cat,
                     "key": _row_block_key(cat, kpi),
                     "prefilled_id": pid,
                     "prefilled_text": patient_text,
                     "prefilled_plan": action_plan})

    # Carry the last non-empty KPI and Categorization labels downward. On a
    # PREFILLED doc the continuation rows have blank kpi/cat (they were merged),
    # so we must rebuild both before grouping, then recompute the block key from
    # the carried-down labels.
    last_kpi = ""
    last_cat = ""
    for d in data:
        if d["kpi"]:
            last_kpi = d["kpi"]
        else:
            d["kpi"] = last_kpi
        if d["cat"]:
            last_cat = d["cat"]
        else:
            d["cat"] = last_cat
        d["key"] = _row_block_key(d["cat"], d["kpi"])

    # Build Categorization groups. Consecutive data rows sharing the same
    # (kpi, cat) belong to one group. Gather each group's prefilled patients
    # (id -> {text, plan}) so we can reconcile them against the Excel.
    groups = []
    cur = None
    for d in data:
        sig = (d["kpi"], d["cat"], d["key"])
        if cur is None or (cur["kpi"], cur["cat"], cur["key"]) != sig:
            cur = {"key": d["key"], "kpi": d["kpi"], "cat": d["cat"],
                   "start": d["idx"], "prefilled": [], "orig_count": 0}
            groups.append(cur)
        cur["orig_count"] += 1
        if d["prefilled_id"] or d["prefilled_text"]:
            cur["prefilled"].append({
                "id": d["prefilled_id"],
                "text": d["prefilled_text"],
                "plan": d["prefilled_plan"],
            })

    # Reconcile each group: produce the final ordered list of (text, plan) rows.
    #   - matched ID  -> keep doc name+reason+plan, update numeric value
    #   - new ID      -> add "Name_ID: value." (no reason/plan)
    #   - stale ID    -> dropped (simply not carried over)
    # Vascular access has no numeric value, so matched rows are kept verbatim.
    for g in groups:
        excl = block_data.get(g["key"], []) if g["key"] else []
        excel_by_id = {}
        excel_order = []
        for (nid, val, bkey) in excl:
            pid = _extract_id(nid)
            if pid and pid not in excel_by_id:
                excel_by_id[pid] = (nid, val, bkey)
                excel_order.append(pid)

        prefilled_by_id = {}
        for pf in g["prefilled"]:
            if pf["id"] and pf["id"] not in prefilled_by_id:
                prefilled_by_id[pf["id"]] = pf

        final_rows = []  # list of {"text":..., "plan":...}
        for pid in excel_order:
            nid, val, bkey = excel_by_id[pid]
            if pid in prefilled_by_id:
                pf = prefilled_by_id[pid]
                # keep existing cell, update the numeric value in place
                new_text = _replace_value_after_id(pf["text"], val, pid)
                final_rows.append({"text": new_text, "plan": pf["plan"]})
            else:
                # new patient: bare entry
                final_rows.append({
                    "text": f"{nid}: {_fmt_value(val, bkey)}.",
                    "plan": "",
                })
        g["rows"] = final_rows

    # Process bottom-up so row insertions don't shift indices above.
    # PHASE 1: resize each group to exactly n rows (add or delete), stamp text.
    total = 0
    for g in reversed(groups):
        start = g["start"]
        frows = g.get("rows", [])
        n = max(1, len(frows))
        orig = g.get("orig_count", 1)

        # Resize the group to n rows.
        if n > orig:
            for _ in range(n - orig):
                insert_row_after(table, start)
        elif n < orig:
            # delete extra rows from the bottom of the group
            for k in range(orig - n):
                delete_row(table, start + n)  # repeatedly remove row at start+n

        cnt = str(len(frows)) if frows else ""
        for i in range(n):
            row = table.rows[start + i]
            set_cell_text(row.cells[0], g["kpi"])
            set_cell_text(row.cells[1], g["cat"])
            set_cell_text(row.cells[2], cnt)
            if i < len(frows):
                set_cell_text(row.cells[3], frows[i]["text"])
                set_cell_text(row.cells[4], frows[i]["plan"])
                total += 1
            else:
                set_cell_text(row.cells[3], "")
                set_cell_text(row.cells[4], "")

    # PHASE 2: apply vertical merges. Read the (now stable) text to compute
    # spans, then merge from the BOTTOM up so continue-cells never need to be
    # re-read by .cells above an unmerged region.
    _apply_merges(table)

    return total


def fill(doc_path, xlsx_path, out_path):
    doc = Document(doc_path)
    n = fill_doc(doc, xlsx_path)
    doc.save(out_path)
    return n


def _data_region(table):
    rows = table.rows
    start = 0
    for ri, row in enumerate(rows):
        joined = (row.cells[0].text + " " + row.cells[1].text).lower()
        if "categorization" in joined:
            start = ri + 1
            break
    end = len(rows)
    for ri in range(start, len(rows)):
        if "attach a report" in rows[ri].cells[0].text.lower():
            end = ri
            break
    return start, end


def _spans(table, col_idx, start, end):
    """Return list of (top, bottom) spans of identical, non-empty labels."""
    rows = table.rows
    out = []
    i = start
    while i < end:
        label = rows[i].cells[col_idx].text.strip()
        if not label:
            i += 1
            continue
        j = i
        while j + 1 < end and rows[j + 1].cells[col_idx].text.strip() == label:
            j += 1
        out.append((i, j))
        i = j + 1
    return out


def _apply_merges(table):
    """Merge KPI (col0), Categorization (col1) and Count (col2) vertically.

    Order matters: compute ALL spans for every column first (while the table is
    still fully unmerged and readable), then apply merges bottom-up.
    """
    start, end = _data_region(table)

    kpi_spans = _spans(table, 0, start, end)
    cat_spans = _spans(table, 1, start, end)
    # Count spans should follow the Categorization grouping (same number across
    # a categorization block), so reuse cat spans for the count column.
    cnt_spans = cat_spans

    # Collect (col, top, bottom, text) then apply bottom-up by top row.
    ops = []
    for (a, b) in kpi_spans:
        if b > a:
            ops.append((0, a, b, table.rows[a].cells[0].text))
    for (a, b) in cat_spans:
        if b > a:
            ops.append((1, a, b, table.rows[a].cells[1].text))
    for (a, b) in cnt_spans:
        if b > a:
            ops.append((2, a, b, table.rows[a].cells[2].text))

    # Apply deepest (largest top index) first so reading .cells for shallower
    # ops is unaffected by continue-merges below them.
    for col, a, b, text in sorted(ops, key=lambda o: -o[1]):
        vmerge_column(table, col, a, b, text=text)
